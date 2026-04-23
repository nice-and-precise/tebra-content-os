#!/usr/bin/env python3
"""
Generate Jordan_Damhof_Resume.pdf from applications/resume.md.

Usage: python3 applications/export/export.py
Requires: pandoc, python-docx, LibreOffice (soffice), pdfinfo
"""

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

REPO = Path(__file__).parent.parent.parent
SRC = REPO / "applications/resume.md"
DOCX = REPO / "applications/export/Jordan_Damhof_Resume.docx"
REF = REPO / "applications/export/compact_reference.docx"
OUT = REPO / "applications/export"


def run(cmd):
    r = subprocess.run(cmd, capture_output=True, text=True)
    if r.returncode:
        print(f"ERROR: {cmd[0]}\n{r.stderr}", file=sys.stderr)
        sys.exit(1)
    return r


def main():
    # 1. pandoc markdown -> docx
    run(["pandoc", str(SRC), "-o", str(DOCX), f"--reference-doc={REF}"])

    # 2. Post-process with python-docx
    doc = Document(str(DOCX))

    # Remove empty HR paragraphs (VML rects from markdown ---)
    hr_paras = [
        p for p in doc.paragraphs
        if not p.text and p._p.find(".//{urn:schemas-microsoft-com:vml}rect") is not None
    ]
    for p in hr_paras:
        p._element.getparent().remove(p._element)

    # Heading 2: tighter size + spacing
    for para in doc.paragraphs:
        if para.style.name == "Heading 2":
            for run in para.runs:
                run.font.size = Pt(12)
            pPr = para._p.get_or_add_pPr()
            sp = pPr.find(qn("w:spacing"))
            if sp is None:
                sp = OxmlElement("w:spacing")
                pPr.append(sp)
            sp.set(qn("w:before"), "60")
            sp.set(qn("w:after"), "20")

    # Exact line spacing on all paragraphs (prevents LibreOffice from adding its own)
    for para in doc.paragraphs:
        pPr = para._p.get_or_add_pPr()
        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            pPr.append(sp)
        sp.set(qn("w:line"), "215")
        sp.set(qn("w:lineRule"), "exact")

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    doc.save(str(DOCX))

    # 3. LibreOffice -> PDF
    run(["soffice", "--headless", "--convert-to", "pdf", str(DOCX), "--outdir", str(OUT)])

    # 4. Verify page count
    r = run(["pdfinfo", str(OUT / "Jordan_Damhof_Resume.pdf")])
    pages = [l for l in r.stdout.splitlines() if l.startswith("Pages:")][0]
    print(pages)
    if "1" not in pages.split()[-1]:
        print("WARNING: PDF is more than 1 page", file=sys.stderr)
    else:
        print("OK: 1-page PDF generated")


if __name__ == "__main__":
    main()
